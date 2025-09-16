#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
SSReportTools - 渗透测试报告生成工具
Python PyQt版本

Author: MaiKeFee
GitHub: https://github.com/Maikefee/
Email: maketoemail@gmail.com
WeChat: rggboom

只有精通计算机理论的人，才具有研究能力。
"""

import sys
import os
import json
import yaml
from datetime import datetime
from pathlib import Path
from PyQt5.QtWidgets import (QApplication, QMainWindow, QWidget, QVBoxLayout, 
                            QHBoxLayout, QGridLayout, QLabel, QLineEdit, 
                            QTextEdit, QComboBox, QPushButton, QTableWidget, 
                            QTableWidgetItem, QTabWidget, QGroupBox, QSpinBox,
                            QDateEdit, QFileDialog, QMessageBox, QSplitter,
                            QHeaderView, QAbstractItemView, QCheckBox)
from PyQt5.QtCore import Qt, QDate, pyqtSignal
from PyQt5.QtGui import QFont, QIcon
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn

class VulnerabilityManager:
    """漏洞库管理器"""
    
    def __init__(self, vuln_file="config/VulnWiki.yml"):
        self.vuln_file = vuln_file
        self.vulnerabilities = {}
        self.load_vulnerabilities()
    
    def load_vulnerabilities(self):
        """加载漏洞库"""
        try:
            with open(self.vuln_file, 'r', encoding='utf-8') as f:
                data = yaml.safe_load(f)
                if data and 'vulnerabilities' in data:
                    for vuln in data['vulnerabilities']:
                        if vuln.get('name'):
                            self.vulnerabilities[vuln['name']] = vuln
        except Exception as e:
            print(f"加载漏洞库失败: {e}")
    
    def get_vulnerability(self, name):
        """获取漏洞信息"""
        return self.vulnerabilities.get(name, {})
    
    def get_all_vulnerabilities(self):
        """获取所有漏洞"""
        return list(self.vulnerabilities.keys())

class TemplateManager:
    """模板管理器"""
    
    def __init__(self, template_dir="config/templates"):
        self.template_dir = template_dir
        self.templates = {}
        self.load_templates()
    
    def load_templates(self):
        """加载模板"""
        template_path = Path(self.template_dir)
        if template_path.exists():
            for json_file in template_path.glob("*.json"):
                try:
                    with open(json_file, 'r', encoding='utf-8') as f:
                        data = json.load(f)
                        template_name = json_file.stem
                        self.templates[template_name] = data
                except Exception as e:
                    print(f"加载模板 {json_file} 失败: {e}")
    
    def get_template(self, name):
        """获取模板"""
        return self.templates.get(name, {})
    
    def get_all_templates(self):
        """获取所有模板"""
        return list(self.templates.keys())

class ReportGenerator:
    """报告生成器"""
    
    def __init__(self, vuln_manager, template_manager):
        self.vuln_manager = vuln_manager
        self.template_manager = template_manager
    
    def generate_report(self, template_name, vuln_data, output_path):
        """生成报告"""
        template = self.template_manager.get_template(template_name)
        if not template:
            raise ValueError(f"模板 {template_name} 不存在")
        
        # 创建Word文档
        doc = Document()
        
        # 设置文档标题
        title = doc.add_heading(f"{template.get('clientName', '')}渗透测试报告", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 添加基本信息
        doc.add_heading('1. 基本信息', level=1)
        
        info_table = doc.add_table(rows=8, cols=2)
        info_table.style = 'Table Grid'
        
        info_data = [
            ('委托单位', template.get('clientName', '')),
            ('测试类型', template.get('isFirstTest', '')),
            ('承测单位', template.get('contractorName', '')),
            ('测试时间', template.get('testDate', '')),
            ('报告日期', f"{template.get('reportYear', '')}年{template.get('reportMonth', '')}月{template.get('reportDay', '')}日"),
            ('报告作者', template.get('reportAuthor', '')),
            ('测试人员', template.get('tester', '')),
            ('项目经理', template.get('manager', ''))
        ]
        
        for i, (key, value) in enumerate(info_data):
            info_table.cell(i, 0).text = key
            info_table.cell(i, 1).text = value
        
        # 添加漏洞统计
        doc.add_heading('2. 漏洞统计', level=1)
        
        stats_table = doc.add_table(rows=2, cols=4)
        stats_table.style = 'Table Grid'
        
        stats_table.cell(0, 0).text = '风险等级'
        stats_table.cell(0, 1).text = '高危'
        stats_table.cell(0, 2).text = '中危'
        stats_table.cell(0, 3).text = '低危'
        
        stats_table.cell(1, 0).text = '数量'
        stats_table.cell(1, 1).text = template.get('highVuln', '0')
        stats_table.cell(1, 2).text = template.get('midVuln', '0')
        stats_table.cell(1, 3).text = template.get('lowVuln', '0')
        
        # 添加漏洞详情
        doc.add_heading('3. 漏洞详情', level=1)
        
        for unit_data in vuln_data:
            unit_name = unit_data.get('unit', '')
            doc.add_heading(f'3.{vuln_data.index(unit_data) + 1} {unit_name}', level=2)
            
            for system_data in unit_data.get('systems', []):
                system_name = system_data.get('system', '')
                doc.add_heading(f'3.{vuln_data.index(unit_data) + 1}.{unit_data.get("systems", []).index(system_data) + 1} {system_name}', level=3)
                
                for vuln in system_data.get('vulns', []):
                    vuln_name = vuln.get('name', '')
                    if vuln_name:
                        vuln_info = self.vuln_manager.get_vulnerability(vuln_name)
                        
                        doc.add_heading(f'3.{vuln_data.index(unit_data) + 1}.{unit_data.get("systems", []).index(system_data) + 1}.{system_data.get("vulns", []).index(vuln) + 1} {vuln_name}', level=4)
                        
                        # 添加漏洞描述
                        if vuln_info.get('description'):
                            doc.add_paragraph(f'漏洞描述：{vuln_info["description"]}')
                        
                        # 添加危害
                        if vuln_info.get('harm'):
                            doc.add_paragraph(f'危害：{vuln_info["harm"]}')
                        
                        # 添加风险等级
                        if vuln_info.get('risklevel'):
                            doc.add_paragraph(f'风险等级：{vuln_info["risklevel"]}')
                        
                        # 添加修复建议
                        if vuln_info.get('suggustion'):
                            doc.add_paragraph(f'修复建议：{vuln_info["suggustion"]}')
                        
                        # 添加修复状态
                        if vuln.get('repaired'):
                            doc.add_paragraph(f'修复状态：{vuln["repaired"]}')
        
        # 保存文档
        doc.save(output_path)
        return output_path

class MainWindow(QMainWindow):
    """主窗口"""
    
    def __init__(self):
        super().__init__()
        self.vuln_manager = VulnerabilityManager()
        self.template_manager = TemplateManager()
        self.report_generator = ReportGenerator(self.vuln_manager, self.template_manager)
        self.init_ui()
    
    def init_ui(self):
        """初始化用户界面"""
        self.setWindowTitle('SSReportTools - 渗透测试报告生成工具 v1.2.0')
        self.setGeometry(100, 100, 1200, 800)
        
        # 显示欢迎对话框
        self.show_welcome_dialog()
        
        # 创建中央widget
        central_widget = QWidget()
        self.setCentralWidget(central_widget)
        
        # 创建主布局
        main_layout = QVBoxLayout(central_widget)
        
        # 创建标签页
        tab_widget = QTabWidget()
        main_layout.addWidget(tab_widget)
        
        # 基本信息标签页
        self.create_basic_info_tab(tab_widget)
        
        # 漏洞管理标签页
        self.create_vulnerability_tab(tab_widget)
        
        # 报告生成标签页
        self.create_report_tab(tab_widget)
        
        # 状态栏
        self.statusBar().showMessage('就绪')
        
        # 添加关于菜单
        menubar = self.menuBar()
        help_menu = menubar.addMenu('帮助')
        
        about_action = help_menu.addAction('关于')
        about_action.triggered.connect(self.show_about)
    
    def create_basic_info_tab(self, parent):
        """创建基本信息标签页"""
        tab = QWidget()
        parent.addTab(tab, '基本信息')
        
        layout = QVBoxLayout(tab)
        
        # 基本信息组
        basic_group = QGroupBox('基本信息')
        basic_layout = QGridLayout(basic_group)
        
        # 基本信息字段
        self.fields = {}
        field_configs = [
            ('clientName', '委托单位', 0, 0),
            ('isFirstTest', '测试类型', 0, 1),
            ('contractorName', '承测单位', 1, 0),
            ('testDate', '测试时间', 1, 1),
            ('reportAuthor', '报告作者', 2, 0),
            ('tester', '测试人员', 2, 1),
            ('manager', '项目经理', 3, 0),
            ('highVuln', '高危漏洞数', 3, 1),
            ('midVuln', '中危漏洞数', 4, 0),
            ('lowVuln', '低危漏洞数', 4, 1)
        ]
        
        for field_name, label, row, col in field_configs:
            label_widget = QLabel(f'{label}:')
            if field_name in ['highVuln', 'midVuln', 'lowVuln']:
                input_widget = QSpinBox()
                input_widget.setRange(0, 999)
            else:
                input_widget = QLineEdit()
            
            basic_layout.addWidget(label_widget, row, col * 2)
            basic_layout.addWidget(input_widget, row, col * 2 + 1)
            self.fields[field_name] = input_widget
        
        # 测试类型下拉框
        test_type_combo = QComboBox()
        test_type_combo.addItems(['初测', '复测'])
        basic_layout.addWidget(QLabel('测试类型:'), 0, 2)
        basic_layout.addWidget(test_type_combo, 0, 3)
        self.fields['isFirstTest'] = test_type_combo
        
        layout.addWidget(basic_group)
        
        # 模板选择组
        template_group = QGroupBox('模板选择')
        template_layout = QHBoxLayout(template_group)
        
        template_layout.addWidget(QLabel('选择模板:'))
        self.template_combo = QComboBox()
        self.template_combo.addItems(self.template_manager.get_all_templates())
        template_layout.addWidget(self.template_combo)
        
        load_template_btn = QPushButton('加载模板')
        load_template_btn.clicked.connect(self.load_template)
        template_layout.addWidget(load_template_btn)
        
        save_template_btn = QPushButton('保存模板')
        save_template_btn.clicked.connect(self.save_template)
        template_layout.addWidget(save_template_btn)
        
        layout.addWidget(template_group)
        layout.addStretch()
    
    def create_vulnerability_tab(self, parent):
        """创建漏洞管理标签页"""
        tab = QWidget()
        parent.addTab(tab, '漏洞管理')
        
        layout = QVBoxLayout(tab)
        
        # 单位管理组
        unit_group = QGroupBox('单位管理')
        unit_layout = QHBoxLayout(unit_group)
        
        unit_layout.addWidget(QLabel('单位名称:'))
        self.unit_name_edit = QLineEdit()
        unit_layout.addWidget(self.unit_name_edit)
        
        add_unit_btn = QPushButton('添加单位')
        add_unit_btn.clicked.connect(self.add_unit)
        unit_layout.addWidget(add_unit_btn)
        
        layout.addWidget(unit_group)
        
        # 系统管理组
        system_group = QGroupBox('系统管理')
        system_layout = QHBoxLayout(system_group)
        
        system_layout.addWidget(QLabel('系统名称:'))
        self.system_name_edit = QLineEdit()
        system_layout.addWidget(self.system_name_edit)
        
        add_system_btn = QPushButton('添加系统')
        add_system_btn.clicked.connect(self.add_system)
        system_layout.addWidget(add_system_btn)
        
        layout.addWidget(system_group)
        
        # 漏洞表格
        self.vuln_table = QTableWidget()
        self.vuln_table.setColumnCount(6)
        self.vuln_table.setHorizontalHeaderLabels(['单位', '系统', '漏洞名称', '风险等级', '修复状态', '操作'])
        
        # 设置表格属性
        self.vuln_table.horizontalHeader().setStretchLastSection(True)
        self.vuln_table.setSelectionBehavior(QAbstractItemView.SelectRows)
        
        layout.addWidget(self.vuln_table)
        
        # 漏洞操作按钮
        vuln_btn_layout = QHBoxLayout()
        
        add_vuln_btn = QPushButton('添加漏洞')
        add_vuln_btn.clicked.connect(self.add_vulnerability)
        vuln_btn_layout.addWidget(add_vuln_btn)
        
        edit_vuln_btn = QPushButton('编辑漏洞')
        edit_vuln_btn.clicked.connect(self.edit_vulnerability)
        vuln_btn_layout.addWidget(edit_vuln_btn)
        
        delete_vuln_btn = QPushButton('删除漏洞')
        delete_vuln_btn.clicked.connect(self.delete_vulnerability)
        vuln_btn_layout.addWidget(delete_vuln_btn)
        
        vuln_btn_layout.addStretch()
        layout.addLayout(vuln_btn_layout)
        
        # 存储漏洞数据
        self.vulnerability_data = []
    
    def create_report_tab(self, parent):
        """创建报告生成标签页"""
        tab = QWidget()
        parent.addTab(tab, '报告生成')
        
        # 创建水平分割器
        splitter = QSplitter(Qt.Horizontal)
        layout = QVBoxLayout(tab)
        layout.addWidget(splitter)
        
        # 左侧：报告生成功能
        left_widget = QWidget()
        left_layout = QVBoxLayout(left_widget)
        
        # 报告生成组
        report_group = QGroupBox('报告生成')
        report_layout = QVBoxLayout(report_group)
        
        # 输出路径选择
        path_layout = QHBoxLayout()
        path_layout.addWidget(QLabel('输出路径:'))
        self.output_path_edit = QLineEdit()
        path_layout.addWidget(self.output_path_edit)
        
        browse_btn = QPushButton('浏览')
        browse_btn.clicked.connect(self.browse_output_path)
        path_layout.addWidget(browse_btn)
        
        report_layout.addLayout(path_layout)
        
        # 生成按钮
        generate_btn = QPushButton('生成报告')
        generate_btn.clicked.connect(self.generate_report)
        generate_btn.setStyleSheet("QPushButton { background-color: #4CAF50; color: white; font-size: 14px; padding: 10px; }")
        report_layout.addWidget(generate_btn)
        
        left_layout.addWidget(report_group)
        
        # 日志显示
        log_group = QGroupBox('生成日志')
        log_layout = QVBoxLayout(log_group)
        
        self.log_text = QTextEdit()
        self.log_text.setReadOnly(True)
        log_layout.addWidget(self.log_text)
        
        left_layout.addWidget(log_group)
        
        # 右侧：作者信息
        right_widget = QWidget()
        right_widget.setMaximumWidth(300)
        right_widget.setStyleSheet("""
            QWidget {
                background-color: #f8f9fa;
                border-left: 2px solid #dee2e6;
            }
        """)
        right_layout = QVBoxLayout(right_widget)
        
        # 作者信息组
        author_group = QGroupBox('作者信息')
        author_group.setStyleSheet("""
            QGroupBox {
                font-weight: bold;
                font-size: 14px;
                color: #2c3e50;
                border: 2px solid #bdc3c7;
                border-radius: 8px;
                margin-top: 10px;
                padding-top: 10px;
            }
            QGroupBox::title {
                subcontrol-origin: margin;
                left: 10px;
                padding: 0 5px 0 5px;
            }
        """)
        author_layout = QVBoxLayout(author_group)
        
        # 作者头像区域
        avatar_label = QLabel('👨‍💻')
        avatar_label.setStyleSheet("font-size: 48px; text-align: center; margin: 10px;")
        avatar_label.setAlignment(Qt.AlignCenter)
        author_layout.addWidget(avatar_label)
        
        # 作者姓名
        name_label = QLabel('MaiKeFee')
        name_label.setStyleSheet("""
            font-size: 18px;
            font-weight: bold;
            color: #2c3e50;
            text-align: center;
            margin: 5px;
        """)
        name_label.setAlignment(Qt.AlignCenter)
        author_layout.addWidget(name_label)
        
        # 个人格言
        motto_label = QLabel('只有精通计算机理论的人，\n才具有研究能力。')
        motto_label.setStyleSheet("""
            font-size: 12px;
            color: #7f8c8d;
            font-style: italic;
            text-align: center;
            margin: 10px;
            padding: 10px;
            background-color: #ecf0f1;
            border-radius: 5px;
        """)
        motto_label.setAlignment(Qt.AlignCenter)
        motto_label.setWordWrap(True)
        author_layout.addWidget(motto_label)
        
        # 分隔线
        line = QLabel()
        line.setStyleSheet("border-top: 1px solid #bdc3c7; margin: 10px 0;")
        author_layout.addWidget(line)
        
        # 联系方式
        contact_label = QLabel('联系方式')
        contact_label.setStyleSheet("""
            font-size: 14px;
            font-weight: bold;
            color: #34495e;
            margin: 5px 0;
        """)
        author_layout.addWidget(contact_label)
        
        # GitHub
        github_layout = QHBoxLayout()
        github_icon = QLabel('🔗')
        github_text = QLabel('GitHub')
        github_text.setStyleSheet("font-weight: bold; color: #2c3e50;")
        github_link = QLabel('<a href="https://github.com/Maikefee/" style="color: #3498db; text-decoration: none;">Maikefee</a>')
        github_link.setOpenExternalLinks(True)
        github_layout.addWidget(github_icon)
        github_layout.addWidget(github_text)
        github_layout.addWidget(github_link)
        github_layout.addStretch()
        author_layout.addLayout(github_layout)
        
        # 邮箱
        email_layout = QHBoxLayout()
        email_icon = QLabel('📧')
        email_text = QLabel('邮箱')
        email_text.setStyleSheet("font-weight: bold; color: #2c3e50;")
        email_value = QLabel('maketoemail@gmail.com')
        email_value.setStyleSheet("color: #7f8c8d;")
        email_layout.addWidget(email_icon)
        email_layout.addWidget(email_text)
        email_layout.addWidget(email_value)
        email_layout.addStretch()
        author_layout.addLayout(email_layout)
        
        # 微信
        wechat_layout = QHBoxLayout()
        wechat_icon = QLabel('💬')
        wechat_text = QLabel('微信')
        wechat_text.setStyleSheet("font-weight: bold; color: #2c3e50;")
        wechat_value = QLabel('rggboom')
        wechat_value.setStyleSheet("color: #7f8c8d;")
        wechat_layout.addWidget(wechat_icon)
        wechat_layout.addWidget(wechat_text)
        wechat_layout.addWidget(wechat_value)
        wechat_layout.addStretch()
        author_layout.addLayout(wechat_layout)
        
        # 分隔线
        line2 = QLabel()
        line2.setStyleSheet("border-top: 1px solid #bdc3c7; margin: 10px 0;")
        author_layout.addWidget(line2)
        
        # 其他项目
        projects_label = QLabel('其他项目')
        projects_label.setStyleSheet("""
            font-size: 14px;
            font-weight: bold;
            color: #34495e;
            margin: 5px 0;
        """)
        author_layout.addWidget(projects_label)
        
        # 盯盘狗量化
        dingpan_layout = QHBoxLayout()
        dingpan_icon = QLabel('🐕')
        dingpan_text = QLabel('盯盘狗量化')
        dingpan_text.setStyleSheet("font-weight: bold; color: #2c3e50;")
        dingpan_link = QLabel('<a href="https://www.dingpandog.com" style="color: #3498db; text-decoration: none;">dingpandog.com</a>')
        dingpan_link.setOpenExternalLinks(True)
        dingpan_layout.addWidget(dingpan_icon)
        dingpan_layout.addWidget(dingpan_text)
        dingpan_layout.addWidget(dingpan_link)
        dingpan_layout.addStretch()
        author_layout.addLayout(dingpan_layout)
        
        # Burp插件
        burp_layout = QHBoxLayout()
        burp_icon = QLabel('🔧')
        burp_text = QLabel('Burp插件')
        burp_text.setStyleSheet("font-weight: bold; color: #2c3e50;")
        burp_value = QLabel('Upload_Auto_Fuzz')
        burp_value.setStyleSheet("color: #7f8c8d; font-size: 10px;")
        burp_layout.addWidget(burp_icon)
        burp_layout.addWidget(burp_text)
        burp_layout.addWidget(burp_value)
        burp_layout.addStretch()
        author_layout.addLayout(burp_layout)
        
        # 量化系统
        quant_layout = QHBoxLayout()
        quant_icon = QLabel('📊')
        quant_text = QLabel('量化系统')
        quant_text.setStyleSheet("font-weight: bold; color: #2c3e50;")
        quant_value = QLabel('Binance_Quantitative')
        quant_value.setStyleSheet("color: #7f8c8d; font-size: 10px;")
        quant_layout.addWidget(quant_icon)
        quant_layout.addWidget(quant_text)
        quant_layout.addWidget(quant_value)
        quant_layout.addStretch()
        author_layout.addLayout(quant_layout)
        
        # 分隔线
        line3 = QLabel()
        line3.setStyleSheet("border-top: 1px solid #bdc3c7; margin: 10px 0;")
        author_layout.addWidget(line3)
        
        # 业务承接
        business_label = QLabel('业务承接')
        business_label.setStyleSheet("""
            font-size: 14px;
            font-weight: bold;
            color: #34495e;
            margin: 5px 0;
        """)
        author_layout.addWidget(business_label)
        
        business_text = QLabel('• 渗透测试\n• 逆向破解\n• Web爬虫\n• 量化策略开发\n• 软件设计\n• 小程序开发')
        business_text.setStyleSheet("""
            color: #7f8c8d;
            font-size: 11px;
            line-height: 1.4;
            margin: 5px;
        """)
        author_layout.addWidget(business_text)
        
        # 添加弹性空间
        author_layout.addStretch()
        
        right_layout.addWidget(author_group)
        
        # 将左右两部分添加到分割器
        splitter.addWidget(left_widget)
        splitter.addWidget(right_widget)
        
        # 设置分割器比例
        splitter.setSizes([800, 300])
    
    def load_template(self):
        """加载模板"""
        template_name = self.template_combo.currentText()
        if not template_name:
            return
        
        template = self.template_manager.get_template(template_name)
        if template:
            # 填充基本信息字段
            for field_name, widget in self.fields.items():
                if field_name in template:
                    if isinstance(widget, QComboBox):
                        widget.setCurrentText(template[field_name])
                    elif isinstance(widget, QSpinBox):
                        widget.setValue(int(template[field_name]) if template[field_name].isdigit() else 0)
                    else:
                        widget.setText(template[field_name])
            
            self.log_message(f"已加载模板: {template_name}")
    
    def save_template(self):
        """保存模板"""
        template_name = self.template_combo.currentText()
        if not template_name:
            QMessageBox.warning(self, '警告', '请选择模板名称')
            return
        
        # 收集基本信息
        template_data = {}
        for field_name, widget in self.fields.items():
            if isinstance(widget, QComboBox):
                template_data[field_name] = widget.currentText()
            elif isinstance(widget, QSpinBox):
                template_data[field_name] = str(widget.value())
            else:
                template_data[field_name] = widget.text()
        
        # 添加当前日期
        now = datetime.now()
        template_data['reportYear'] = str(now.year)
        template_data['reportMonth'] = f"{now.month:02d}"
        template_data['reportDay'] = f"{now.day:02d}"
        
        # 保存到文件
        template_path = Path(self.template_manager.template_dir) / f"{template_name}.json"
        try:
            with open(template_path, 'w', encoding='utf-8') as f:
                json.dump(template_data, f, ensure_ascii=False, indent=2)
            self.log_message(f"模板已保存: {template_path}")
            QMessageBox.information(self, '成功', '模板保存成功')
        except Exception as e:
            QMessageBox.critical(self, '错误', f'保存模板失败: {e}')
    
    def add_unit(self):
        """添加单位"""
        unit_name = self.unit_name_edit.text().strip()
        if not unit_name:
            QMessageBox.warning(self, '警告', '请输入单位名称')
            return
        
        # 检查是否已存在
        for unit in self.vulnerability_data:
            if unit['unit'] == unit_name:
                QMessageBox.warning(self, '警告', '该单位已存在')
                return
        
        # 添加新单位
        self.vulnerability_data.append({
            'unit': unit_name,
            'systems': []
        })
        
        self.unit_name_edit.clear()
        self.log_message(f"已添加单位: {unit_name}")
    
    def add_system(self):
        """添加系统"""
        system_name = self.system_name_edit.text().strip()
        if not system_name:
            QMessageBox.warning(self, '警告', '请输入系统名称')
            return
        
        # 添加到最后一个单位
        if not self.vulnerability_data:
            QMessageBox.warning(self, '警告', '请先添加单位')
            return
        
        last_unit = self.vulnerability_data[-1]
        last_unit['systems'].append({
            'system': system_name,
            'vulns': []
        })
        
        self.system_name_edit.clear()
        self.log_message(f"已添加系统: {system_name}")
    
    def add_vulnerability(self):
        """添加漏洞"""
        if not self.vulnerability_data:
            QMessageBox.warning(self, '警告', '请先添加单位和系统')
            return
        
        # 获取可用的漏洞列表
        available_vulns = self.vuln_manager.get_all_vulnerabilities()
        if not available_vulns:
            QMessageBox.warning(self, '警告', '没有可用的漏洞库')
            return
        
        # 创建添加漏洞对话框
        from PyQt5.QtWidgets import QDialog, QVBoxLayout, QLabel, QComboBox, QPushButton, QHBoxLayout
        
        dialog = QDialog(self)
        dialog.setWindowTitle('添加漏洞')
        dialog.setModal(True)
        dialog.resize(500, 300)
        
        layout = QVBoxLayout(dialog)
        
        # 选择单位
        layout.addWidget(QLabel('选择单位:'))
        unit_combo = QComboBox()
        for unit in self.vulnerability_data:
            unit_combo.addItem(unit['unit'])
        layout.addWidget(unit_combo)
        
        # 选择系统
        layout.addWidget(QLabel('选择系统:'))
        system_combo = QComboBox()
        def update_systems():
            system_combo.clear()
            selected_unit = unit_combo.currentText()
            for unit in self.vulnerability_data:
                if unit['unit'] == selected_unit:
                    for system in unit['systems']:
                        system_combo.addItem(system['system'])
                    break
        unit_combo.currentTextChanged.connect(update_systems)
        update_systems()  # 初始化
        layout.addWidget(system_combo)
        
        # 选择漏洞类型
        layout.addWidget(QLabel('选择漏洞类型:'))
        vuln_combo = QComboBox()
        vuln_combo.addItems(available_vulns)
        layout.addWidget(vuln_combo)
        
        # 修复状态
        layout.addWidget(QLabel('修复状态:'))
        repaired_combo = QComboBox()
        repaired_combo.addItems(['未修复', '已修复', '修复中', '不适用'])
        repaired_combo.setCurrentText('未修复')
        layout.addWidget(repaired_combo)
        
        # 风险等级
        layout.addWidget(QLabel('风险等级:'))
        risk_combo = QComboBox()
        risk_combo.addItems(['高危', '中危', '低危', '信息'])
        risk_combo.setCurrentText('高危')
        layout.addWidget(risk_combo)
        
        # 按钮
        button_layout = QHBoxLayout()
        add_btn = QPushButton('添加')
        add_btn.clicked.connect(dialog.accept)
        cancel_btn = QPushButton('取消')
        cancel_btn.clicked.connect(dialog.reject)
        button_layout.addWidget(add_btn)
        button_layout.addWidget(cancel_btn)
        layout.addLayout(button_layout)
        
        if dialog.exec_() == QDialog.Accepted:
            unit_name = unit_combo.currentText()
            system_name = system_combo.currentText()
            vuln_name = vuln_combo.currentText()
            repaired = repaired_combo.currentText()
            risk_level = risk_combo.currentText()
            
            # 添加到数据源
            for unit in self.vulnerability_data:
                if unit['unit'] == unit_name:
                    for system in unit['systems']:
                        if system['system'] == system_name:
                            system['vulns'].append({
                                'name': vuln_name,
                                'repaired': repaired,
                                'risk_level': risk_level
                            })
                            break
                    break
            
            self.update_vulnerability_table()
            self.log_message(f"已添加漏洞: {vuln_name} - 单位: {unit_name}, 系统: {system_name}")
            QMessageBox.information(self, '成功', '漏洞添加成功')
    
    def edit_vulnerability(self):
        """编辑漏洞"""
        current_row = self.vuln_table.currentRow()
        if current_row < 0:
            QMessageBox.warning(self, '警告', '请选择要编辑的漏洞')
            return
        
        # 获取当前选中的漏洞信息
        unit_name = self.vuln_table.item(current_row, 0).text()
        system_name = self.vuln_table.item(current_row, 1).text()
        vuln_name = self.vuln_table.item(current_row, 2).text()
        current_repaired = self.vuln_table.item(current_row, 4).text()
        
        # 创建编辑对话框
        from PyQt5.QtWidgets import QDialog, QFormLayout, QComboBox, QDialogButtonBox
        
        dialog = QDialog(self)
        dialog.setWindowTitle('编辑漏洞')
        dialog.setModal(True)
        dialog.resize(400, 200)
        
        layout = QFormLayout(dialog)
        
        # 漏洞名称（只读）
        vuln_name_label = QLabel(vuln_name)
        vuln_name_label.setStyleSheet("font-weight: bold; color: #2c3e50;")
        layout.addRow('漏洞名称:', vuln_name_label)
        
        # 修复状态
        repaired_combo = QComboBox()
        repaired_combo.addItems(['未修复', '已修复', '修复中', '不适用'])
        repaired_combo.setCurrentText(current_repaired)
        layout.addRow('修复状态:', repaired_combo)
        
        # 风险等级
        risk_combo = QComboBox()
        risk_combo.addItems(['高危', '中危', '低危', '信息'])
        current_risk = self.vuln_table.item(current_row, 3).text()
        risk_combo.setCurrentText(current_risk if current_risk else '高危')
        layout.addRow('风险等级:', risk_combo)
        
        # 按钮
        buttons = QDialogButtonBox(QDialogButtonBox.Ok | QDialogButtonBox.Cancel)
        buttons.accepted.connect(dialog.accept)
        buttons.rejected.connect(dialog.reject)
        layout.addRow(buttons)
        
        if dialog.exec_() == QDialog.Accepted:
            # 更新漏洞信息
            new_repaired = repaired_combo.currentText()
            new_risk = risk_combo.currentText()
            
            # 更新表格显示
            self.vuln_table.setItem(current_row, 3, QTableWidgetItem(new_risk))
            self.vuln_table.setItem(current_row, 4, QTableWidgetItem(new_repaired))
            
            # 更新数据源
            self.update_vulnerability_data(unit_name, system_name, vuln_name, new_repaired, new_risk)
            
            self.log_message(f"已更新漏洞: {vuln_name} - 状态: {new_repaired}, 风险: {new_risk}")
            QMessageBox.information(self, '成功', '漏洞信息已更新')
    
    def update_vulnerability_data(self, unit_name, system_name, vuln_name, new_repaired, new_risk):
        """更新漏洞数据源"""
        for unit in self.vulnerability_data:
            if unit['unit'] == unit_name:
                for system in unit['systems']:
                    if system['system'] == system_name:
                        for vuln in system['vulns']:
                            if vuln['name'] == vuln_name:
                                vuln['repaired'] = new_repaired
                                vuln['risk_level'] = new_risk
                                return
    
    def delete_vulnerability(self):
        """删除漏洞"""
        current_row = self.vuln_table.currentRow()
        if current_row < 0:
            QMessageBox.warning(self, '警告', '请选择要删除的漏洞')
            return
        
        # 获取要删除的漏洞信息
        unit_name = self.vuln_table.item(current_row, 0).text()
        system_name = self.vuln_table.item(current_row, 1).text()
        vuln_name = self.vuln_table.item(current_row, 2).text()
        
        # 确认删除
        reply = QMessageBox.question(
            self, 
            '确认删除', 
            f'确定要删除漏洞 "{vuln_name}" 吗？\n\n单位: {unit_name}\n系统: {system_name}',
            QMessageBox.Yes | QMessageBox.No,
            QMessageBox.No
        )
        
        if reply == QMessageBox.Yes:
            # 从数据源中删除
            if self.remove_vulnerability_from_data(unit_name, system_name, vuln_name):
                # 从表格中删除行
                self.vuln_table.removeRow(current_row)
                self.log_message(f"已删除漏洞: {vuln_name}")
                QMessageBox.information(self, '成功', '漏洞已删除')
            else:
                QMessageBox.warning(self, '错误', '删除漏洞失败')
    
    def remove_vulnerability_from_data(self, unit_name, system_name, vuln_name):
        """从数据源中删除漏洞"""
        for unit in self.vulnerability_data:
            if unit['unit'] == unit_name:
                for system in unit['systems']:
                    if system['system'] == system_name:
                        for i, vuln in enumerate(system['vulns']):
                            if vuln['name'] == vuln_name:
                                system['vulns'].pop(i)
                                return True
        return False
    
    def update_vulnerability_table(self):
        """更新漏洞表格"""
        self.vuln_table.setRowCount(0)
        
        for unit in self.vulnerability_data:
            for system in unit['systems']:
                for vuln in system['vulns']:
                    row = self.vuln_table.rowCount()
                    self.vuln_table.insertRow(row)
                    
                    self.vuln_table.setItem(row, 0, QTableWidgetItem(unit['unit']))
                    self.vuln_table.setItem(row, 1, QTableWidgetItem(system['system']))
                    self.vuln_table.setItem(row, 2, QTableWidgetItem(vuln['name']))
                    
                    # 获取风险等级（优先使用用户设置的值）
                    risk_level = vuln.get('risk_level')
                    if not risk_level:
                        # 如果没有用户设置的值，从漏洞库获取
                        vuln_info = self.vuln_manager.get_vulnerability(vuln['name'])
                        risk_level = vuln_info.get('risklevel', '未知')
                    self.vuln_table.setItem(row, 3, QTableWidgetItem(risk_level))
                    self.vuln_table.setItem(row, 4, QTableWidgetItem(vuln.get('repaired', '未修复')))
                    
                    # 操作按钮
                    edit_btn = QPushButton('编辑')
                    edit_btn.clicked.connect(lambda checked, row=row: self.edit_vulnerability())
                    self.vuln_table.setCellWidget(row, 5, edit_btn)
    
    def browse_output_path(self):
        """浏览输出路径"""
        file_path, _ = QFileDialog.getSaveFileName(
            self, '保存报告', '', 'Word文档 (*.docx)'
        )
        if file_path:
            self.output_path_edit.setText(file_path)
    
    def generate_report(self):
        """生成报告"""
        if not self.output_path_edit.text():
            QMessageBox.warning(self, '警告', '请选择输出路径')
            return
        
        if not self.vulnerability_data:
            QMessageBox.warning(self, '警告', '请添加漏洞数据')
            return
        
        try:
            template_name = self.template_combo.currentText()
            output_path = self.output_path_edit.text()
            
            self.log_message("开始生成报告...")
            
            # 生成报告
            result_path = self.report_generator.generate_report(
                template_name, 
                self.vulnerability_data, 
                output_path
            )
            
            self.log_message(f"报告生成成功: {result_path}")
            QMessageBox.information(self, '成功', f'报告已生成: {result_path}')
            
        except Exception as e:
            self.log_message(f"生成报告失败: {e}")
            QMessageBox.critical(self, '错误', f'生成报告失败: {e}')
    
    def log_message(self, message):
        """添加日志消息"""
        timestamp = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        self.log_text.append(f"[{timestamp}] {message}")
    
    def show_about(self):
        """显示关于对话框"""
        from PyQt5.QtWidgets import QDialog, QVBoxLayout, QLabel, QPushButton, QHBoxLayout
        from PyQt5.QtCore import Qt
        
        dialog = QDialog(self)
        dialog.setWindowTitle('关于 SSReportTools')
        dialog.setModal(True)
        dialog.resize(500, 400)
        
        layout = QVBoxLayout(dialog)
        
        # 标题
        title_label = QLabel('SSReportTools')
        title_label.setStyleSheet("font-size: 24px; font-weight: bold; color: #2c3e50; margin: 10px;")
        title_label.setAlignment(Qt.AlignCenter)
        layout.addWidget(title_label)
        
        # 版本信息
        version_label = QLabel('版本: v1.2.0')
        version_label.setStyleSheet("font-size: 14px; color: #7f8c8d; margin: 5px;")
        version_label.setAlignment(Qt.AlignCenter)
        layout.addWidget(version_label)
        
        # 描述
        desc_label = QLabel('渗透测试报告生成工具\n基于Python PyQt5开发')
        desc_label.setStyleSheet("font-size: 12px; color: #34495e; margin: 10px;")
        desc_label.setAlignment(Qt.AlignCenter)
        layout.addWidget(desc_label)
        
        # 作者信息
        author_label = QLabel('作者: MaiKeFee')
        author_label.setStyleSheet("font-size: 14px; font-weight: bold; color: #2c3e50; margin: 10px;")
        author_label.setAlignment(Qt.AlignCenter)
        layout.addWidget(author_label)
        
        # 联系方式
        contact_text = """
GitHub: https://github.com/Maikefee/
Email: maketoemail@gmail.com
WeChat: rggboom

只有精通计算机理论的人，才具有研究能力。
        """
        contact_label = QLabel(contact_text)
        contact_label.setStyleSheet("font-size: 11px; color: #7f8c8d; margin: 10px;")
        contact_label.setAlignment(Qt.AlignCenter)
        layout.addWidget(contact_label)
        
        # 功能特性
        features_text = """
功能特性:
• 专业漏洞库 (30+ 漏洞类型)
• OWASP Top 10 2021 支持
• 完整漏洞管理 (增删改查)
• 智能报告生成
• 跨平台支持
        """
        features_label = QLabel(features_text)
        features_label.setStyleSheet("font-size: 11px; color: #34495e; margin: 10px;")
        features_label.setAlignment(Qt.AlignLeft)
        layout.addWidget(features_label)
        
        # 按钮
        button_layout = QHBoxLayout()
        close_btn = QPushButton('关闭')
        close_btn.clicked.connect(dialog.accept)
        close_btn.setStyleSheet("QPushButton { background-color: #3498db; color: white; padding: 8px 16px; border-radius: 4px; }")
        button_layout.addWidget(close_btn)
        layout.addLayout(button_layout)
        
        dialog.exec_()
    
    def show_welcome_dialog(self):
        """显示欢迎对话框"""
        from PyQt5.QtWidgets import QDialog, QVBoxLayout, QLabel, QPushButton, QHBoxLayout
        from PyQt5.QtCore import Qt, QTimer
        
        dialog = QDialog(self)
        dialog.setWindowTitle('欢迎使用 SSReportTools')
        dialog.setModal(True)
        dialog.resize(600, 500)
        dialog.setStyleSheet("""
            QDialog {
                background-color: #f8f9fa;
            }
            QLabel {
                color: #2c3e50;
            }
        """)
        
        layout = QVBoxLayout(dialog)
        
        # 标题
        title_label = QLabel('SSReportTools')
        title_label.setStyleSheet("""
            font-size: 32px; 
            font-weight: bold; 
            color: #2c3e50; 
            margin: 20px;
            background: qlineargradient(x1:0, y1:0, x2:1, y2:0, stop:0 #3498db, stop:1 #2980b9);
            color: white;
            padding: 15px;
            border-radius: 10px;
        """)
        title_label.setAlignment(Qt.AlignCenter)
        layout.addWidget(title_label)
        
        # 版本信息
        version_label = QLabel('渗透测试报告生成工具 v1.2.0')
        version_label.setStyleSheet("font-size: 16px; color: #7f8c8d; margin: 10px;")
        version_label.setAlignment(Qt.AlignCenter)
        layout.addWidget(version_label)
        
        # 作者信息
        author_section = QLabel()
        author_section.setText("""
        <div style='background-color: #ecf0f1; padding: 15px; border-radius: 8px; margin: 10px;'>
            <h3 style='color: #2c3e50; margin: 0 0 10px 0;'>👨‍💻 作者信息</h3>
            <p style='margin: 5px 0;'><strong>作者:</strong> MaiKeFee</p>
            <p style='margin: 5px 0;'><strong>GitHub:</strong> <a href='https://github.com/Maikefee/' style='color: #3498db;'>https://github.com/Maikefee/</a></p>
            <p style='margin: 5px 0;'><strong>邮箱:</strong> maketoemail@gmail.com</p>
            <p style='margin: 5px 0;'><strong>微信:</strong> rggboom</p>
            <p style='margin: 10px 0 5px 0; font-style: italic; color: #7f8c8d;'>💡 只有精通计算机理论的人，才具有研究能力。</p>
        </div>
        """)
        author_section.setWordWrap(True)
        layout.addWidget(author_section)
        
        # 功能特性
        features_section = QLabel()
        features_section.setText("""
        <div style='background-color: #e8f5e8; padding: 15px; border-radius: 8px; margin: 10px;'>
            <h3 style='color: #27ae60; margin: 0 0 10px 0;'>🚀 功能特性</h3>
            <ul style='margin: 0; padding-left: 20px;'>
                <li>🔒 专业漏洞库 (30+ 漏洞类型)</li>
                <li>📊 OWASP Top 10 2021 支持</li>
                <li>🛠️ 完整漏洞管理 (增删改查)</li>
                <li>📝 智能报告生成</li>
                <li>🌐 跨平台支持</li>
            </ul>
        </div>
        """)
        features_section.setWordWrap(True)
        layout.addWidget(features_section)
        
        # 其他项目
        projects_section = QLabel()
        projects_section.setText("""
        <div style='background-color: #fff3cd; padding: 15px; border-radius: 8px; margin: 10px;'>
            <h3 style='color: #856404; margin: 0 0 10px 0;'>🌟 其他项目</h3>
            <p style='margin: 5px 0;'><strong>盯盘狗量化:</strong> <a href='https://www.dingpandog.com' style='color: #3498db;'>https://www.dingpandog.com</a></p>
            <p style='margin: 5px 0;'><strong>Burp插件:</strong> Upload_Auto_Fuzz_Burp_Plugin, Fofa_Google_Plugin</p>
            <p style='margin: 5px 0;'><strong>量化系统:</strong> Binance_Quantitative_System</p>
        </div>
        """)
        projects_section.setWordWrap(True)
        layout.addWidget(projects_section)
        
        # 按钮
        button_layout = QHBoxLayout()
        
        # 开始使用按钮
        start_btn = QPushButton('🚀 开始使用')
        start_btn.setStyleSheet("""
            QPushButton {
                background-color: #3498db;
                color: white;
                font-size: 14px;
                font-weight: bold;
                padding: 12px 24px;
                border-radius: 6px;
                border: none;
            }
            QPushButton:hover {
                background-color: #2980b9;
            }
            QPushButton:pressed {
                background-color: #21618c;
            }
        """)
        start_btn.clicked.connect(dialog.accept)
        
        # 关于按钮
        about_btn = QPushButton('ℹ️ 关于')
        about_btn.setStyleSheet("""
            QPushButton {
                background-color: #95a5a6;
                color: white;
                font-size: 14px;
                padding: 12px 24px;
                border-radius: 6px;
                border: none;
            }
            QPushButton:hover {
                background-color: #7f8c8d;
            }
        """)
        about_btn.clicked.connect(lambda: (dialog.accept(), self.show_about()))
        
        button_layout.addWidget(start_btn)
        button_layout.addWidget(about_btn)
        layout.addLayout(button_layout)
        
        # 自动关闭定时器（可选）
        timer = QTimer()
        timer.timeout.connect(dialog.accept)
        timer.setSingleShot(True)
        timer.start(10000)  # 10秒后自动关闭
        
        dialog.exec_()

def main():
    """主函数"""
    app = QApplication(sys.argv)
    
    # 设置应用程序信息
    app.setApplicationName('SSReportTools')
    app.setApplicationVersion('1.0.0')
    app.setOrganizationName('Security Tools')
    
    # 创建主窗口
    window = MainWindow()
    window.show()
    
    sys.exit(app.exec_())

if __name__ == '__main__':
    main()
